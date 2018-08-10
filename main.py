# coding:utf-8
import os
import re
import urllib.request

import shutil

import docx
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Cm, RGBColor

url = "http://222.197.183.99/TutorList.aspx"
pattern = "div"
html_class = "NewsDetails"
domain = "http://222.197.183.99"

key_words = ["导师代码", "导师姓名", "性别", "出生年月", "特称", "职称", "学位", "属性", "电子邮件",
             "学术经历", "个人简介", "科研项目", "发表文章", "博士招生专业", "硕士招生专业"]

def instructor_pageaddr_list(html_str, school_name):
    inst_soup = BeautifulSoup(html_str[html_str.find(school_name):], "lxml")
    addrs_list = {}
    for link_div in inst_soup.find("td").find_all("div"):
        url = link_div.find("a").get("href")
        name = re.sub("\d+\s", "", link_div.find("a").get_text().strip())
        url = url if url[0] != "/" else url[1:]
        addrs_list[name] = domain + "/" + url
    return addrs_list


def schools_list(html_str):
    schools = {}
    all_list_soup = BeautifulSoup(html_str, "lxml")
    for school in all_list_soup.find_all("b"):
        if re.match("\d{3}", school.get_text()):
            schools[int(school.get_text()[:3])] = school.get_text()[3:].strip()
    return schools


def construct_instructor_info(url):
    response = None
    for i in range(2, 6):
        try:
            response = urllib.request.urlopen(url, timeout=10)
            break
        except Exception as e:
            print(e, "获取导师信息超时, 尝试第" + str(i) + "次连接: " + url)

    if response is None:
        return None

    html = response.read().decode(encoding="utf8", errors="ignore")
    html_soup = BeautifulSoup(str(html).replace("&nbsp;", ""), "lxml").find("table")

    instructor = {}

    instructor["网页地址"] = url
    instructor["学院"] = html_soup.find("span", attrs={"id": "Labelxymc"}).get_text().strip()
    instructor["照片"] = domain + html_soup.find("img").get("src").strip()[1:]

    html = re.sub("\s{2,}", "", str(html_soup))
    html = re.sub("<style((.*)\n)+</style>", "", html)

    pos = 0
    for key_word in key_words:
        pos = html.find(key_word, pos)
        soup = BeautifulSoup(html[pos:], "lxml")
        instructor[key_word] = soup.find("span").get_text().strip()

    table_soup = BeautifulSoup(html, "lxml").find("table", attrs={"cellspacing": "0", "cellpadding": "0", "width": "100%", "border": "0"})

    '''
    tuple[0]两个td标签, 一个博士专业代码及名称, 一个硕士专业代码及名称
    tuple[1]两个table标签分别嵌套在td中, 一个博士方向, 一个硕士方向, 如果没有方向, 则没有table
    '''
    pos_mark = [(2, 3), (5, 6), (8, 9), (11, 12), (14, 15)]
    trs = table_soup.find_all("tr", recursive=False)

    bs_spciality_code_label = "Labelbszydm"
    bs_spciality_name_label = "Labelbszymc"
    ss_spciality_code_label = "Labelsszydm"
    ss_spciality_name_label = "Labelsszymc"

    bs_spcialitys = []
    ss_spcialitys = []

    for i in range(len(pos_mark)):
        speciality = trs[pos_mark[i][0]]
        # 获取单个条目的专业名称(包含代码)
        # 博士专业
        bs_spciality_name = speciality.find("span", attrs={"id": bs_spciality_code_label + str(i + 1)}).get_text() + " " + speciality.find("span", attrs={"id": bs_spciality_name_label + str(i + 1)}).get_text()
        bs_spciality_name = bs_spciality_name.strip()
        # 硕士专业
        ss_spciality_name = speciality.find("span", attrs={"id": ss_spciality_code_label + str(i + 1)}).get_text() + " " + speciality.find("span", attrs={"id": ss_spciality_name_label + str(i + 1)}).get_text()
        ss_spciality_name = ss_spciality_name.strip()


        # 获取博士和硕士的专业方向, html分别放在table1, table2
        tds = trs[pos_mark[i][1]].find_all("td", recursive=False)
        table1 = tds[0].find("table")
        table2 = tds[1].find("table")

        single_bs_speciality = {}
        single_ss_speciality = {}

        bs_orientations = []
        ss_orientations = []

        if table1:
            for tr in table1.find_all("tr", recursive=False):
                bs_orientations.append(tr.find("td", attrs={"class": "width4em"}).get_text().strip() +
                                      tr.find("td", attrs={"class": "alignleft"}).get_text().strip())

        if table2:
            for tr in table2.find_all("tr", recursive=False):
                ss_orientations.append(tr.find("td", attrs={"class": "width4em"}).get_text().strip() +
                                      tr.find("td", attrs={"class": "alignleft"}).get_text().strip())

        if len(bs_spciality_name.strip()) != 0:
            single_bs_speciality[bs_spciality_name] = bs_orientations
            bs_spcialitys.append(single_bs_speciality)

        if len(ss_spciality_name.strip()) != 0:
            single_ss_speciality[ss_spciality_name] = ss_orientations
            ss_spcialitys.append(single_ss_speciality)

    instructor["博士招生专业"] = bs_spcialitys
    instructor["硕士招生专业"] = ss_spcialitys

    return instructor


def show_instructor(instructor, file):
    for key_word in ["网页地址", "照片", "学院"] + key_words:
        if (key_word == "博士招生专业") or (key_word == "硕士招生专业"):
            if len(instructor[key_word]) == 0:
                # print("\n[" + key_word + "]: 不招生", end="")
                file.write("\n[" + key_word + "]: 不招生\n")
            else:
                # print("\n[" + key_word + "]: ")
                file.write("\n[" + key_word + "]: \n")

            for speciality in instructor[key_word]: # 专业到方向列表的映射
                for x in speciality:
                    # print("\t(" + x +"): ")
                    file.write("\t(" + x +"): \n")
                    for oriention in speciality[x]:
                        # print("\t\t\t", oriention)
                        file.write("\t\t\t" + oriention + "\n")
            # print()
            file.write("\n")
        elif key_word == "学术经历" or key_word == "个人简介" or key_word == "科研项目" or key_word == "发表文章":
            if len(instructor[key_word]) == 0:
                # print("[" + key_word + "]: 未介绍")
                file.write("[" + key_word + "]: 未介绍\n")
            else:
                # print("\n[" + key_word + "]: ")
                # print("\t\t", instructor[key_word])
                # print()
                if key_word == "发表文章" and len(instructor["发表文章"]) > 500:
                    file.write("\n[" + key_word + "]: \n" + "\t\t" + "该项太长, 请到网页中浏览" + "\n\n")
                else:
                    file.write("\n[" + key_word + "]: \n" + "\t\t" + instructor[key_word] + "\n\n")
        else:
            wrap_word = key_word if len(key_word) == 4 else "      " + key_word
            # print("%6s: %s" % (wrap_word, instructor[key_word] if len(instructor[key_word]) != 0 else "未介绍"))
            file.write("%6s: %s\n" % (wrap_word, instructor[key_word] if len(instructor[key_word]) != 0 else "未介绍"))

    file.write("===============================================================================================\n\n\n")
    # print("===============================================================================================\n")


def gen_speciality_str(specialitys):
    speciality_str = ""
    if len(specialitys) == 0:
        speciality_str = "博士不招生"
    else:
        for speciality in specialitys:  # 专业到方向列表的映射
            for x in speciality:
                speciality_str += "[" + x + "]:"
                for oriention in speciality[x]:
                    speciality_str += ("\n\t" + oriention)
                speciality_str += "\n"
    return speciality_str


def write_docx(instructor, sample_docx, dir_name):
    document = Document(sample_docx)
    outer_table = document.tables[0]
    inner_table = outer_table.cell(0, 1).tables[0]

    pic_file_name = instructor["照片"][instructor["照片"].rfind("/") + 1:]
    for i in range(2, 6):
        try:
            response = urllib.request.urlopen(instructor["照片"], timeout=10)
            f = open(pic_file_name, "w+b")
            f.write(response.read())
            f.close()
            break
        except Exception as e:
            print(e, "获取头像超时, 尝试第" + str(i) + "次连接: " + instructor["照片"])

    paragraph = outer_table.rows[0].cells[0].paragraphs[0]
    run = paragraph.add_run()

    if os.path.exists(pic_file_name) == True:
        # run.add_picture(temp_pic_file, width=Cm(4.2), height=Cm(6.2)) # 窄边框使用
        try:
            run.add_picture(pic_file_name, width=Cm(4.0), height=Cm(5.7)) # 厚边框使用
            os.remove(pic_file_name)
        except Exception as e:
            run.add_picture("default.gif", width=Cm(4.0), height=Cm(5.7))  # 厚边框使用
    else:
        # run.add_picture(temp_pic_file, width=Cm(4.2), height=Cm(6.2)) # 窄边框使用
        run.add_picture("default.gif", width=Cm(4.0), height=Cm(5.7))  # 厚边框使用

    inner_table.cell(0, 1).text = instructor["导师代码"]
    inner_table.cell(1, 1).text = instructor["导师姓名"]
    inner_table.cell(2, 1).text = instructor["性别"]
    inner_table.cell(3, 1).text = instructor["出生年月"]
    inner_table.cell(4, 1).text = instructor["特称"] if (len(instructor["特称"])) != 0 else "未介绍"
    inner_table.cell(5, 1).text = instructor["职称"] if (len(instructor["职称"])) != 0 else "未介绍"
    inner_table.cell(6, 1).text = instructor["学位"] if (len(instructor["学位"])) != 0 else "未介绍"
    inner_table.cell(7, 1).text = instructor["属性"] if (len(instructor["属性"])) != 0 else "未介绍"
    inner_table.cell(8, 1).text = instructor["电子邮件"]
    inner_table.cell(9, 1).text = instructor["学院"]
    inner_table.cell(10, 1).text = instructor["网页地址"]

    paragraphs = [p for p in document.paragraphs]

    # 添加导师姓名在最开头并且是一级标题
    run = paragraphs[0].add_run()
    font = run.font
    run.add_text(instructor["导师姓名"])
    font.color.rgb = RGBColor(255, 0, 0)

    # 分别添加导师的每一项信息
    for i in range(len(paragraphs)):
        if paragraphs[i].text == "个人简介":
            document.paragraphs[i + 1].text = instructor["个人简介"] if len(instructor["个人简介"]) != 0 else "未介绍"
        elif paragraphs[i].text == "学术经历":
            document.paragraphs[i + 1].text = instructor["学术经历"] if len(instructor["学术经历"]) != 0 else "未介绍"
        elif paragraphs[i].text == "科研项目":
            document.paragraphs[i + 1].text = instructor["科研项目"] if len(instructor["科研项目"]) != 0 else "未介绍"
        elif paragraphs[i].text == "发表文章":
            if len(instructor["发表文章"]) <= 500:
                document.paragraphs[i + 1].text = instructor["发表文章"] if len(instructor["发表文章"]) != 0 else "未介绍"
            else:
                document.paragraphs[i + 1].text = "该导师发表文章项内容超级长, 为了不影响参阅, 详见请打开导师介绍页面"
        elif paragraphs[i].text == "博士招生专业":
            document.paragraphs[i + 1].text = gen_speciality_str(instructor["博士招生专业"])
        elif paragraphs[i].text == "硕士招生专业":
            document.paragraphs[i + 1].text = gen_speciality_str(instructor["硕士招生专业"])

    document.add_page_break()
    document.save(dir_name + "/" + "[" + instructor["导师代码"] + "]" + instructor["导师姓名"] + ".docx")


def debug():
    ins_url = "http://222.197.183.99/TutorDetails.aspx?id=308"
    #ins_url = "http://222.197.183.99/TutorDetails.aspx?id=2257" #出现以为HTML标签
    ins_url = "http://222.197.183.99/TutorDetails.aspx?id=113"
    ins_url = "http://222.197.183.99/TutorDetails.aspx?id=2177"
    instructor = construct_instructor_info(ins_url)
    write_docx(instructor, "sample.docx", ".")
    show_instructor(instructor, open("uestc.txt", "w"))
    exit(0)

if 1 == 91:
    debug()


# ================================start==================================== #
response = None
for i in range(2, 6):
    try:
        response = urllib.request.urlopen(url, timeout=10)
        break
    except Exception as e:
        print(e, "网络访问超时, 尝试第" + str(i) + "次连接: " + url)

if response is None:
    print("网络访问失败")
    exit(0)

html = response.read().decode(encoding="utf8", errors="ignore")


full_page_soup = BeautifulSoup(str(html), "lxml")
all_list_soup = full_page_soup.body.find(pattern, attrs = {"class": html_class})

schools = schools_list(str(all_list_soup))
for x in schools.items():
    print(x)

school_id = input("输入学院ID:")
list_url = instructor_pageaddr_list(str(all_list_soup), schools[int(school_id)])

dir_name = "generatefiles"
sample_docx = "sample.docx"

if os.path.exists(dir_name):
    shutil.rmtree(dir_name)
os.mkdir(dir_name)

f = open(dir_name + "/" + "导师列表纯文本.txt", "w")

for index, x in enumerate(list_url.items()):
    print("抓取页面: " + x[0] + " " + x[1])
    instructor_info = construct_instructor_info(x[1])
    if instructor_info is not None:
        show_instructor(instructor_info, f)
        write_docx(instructor_info, sample_docx, dir_name)
    else:
        print("获取导师信息失败:", x)
    print(index + 1, "完成创建: " + "[" + instructor_info["导师代码"] + "]" + instructor_info["导师姓名"] + ".docx\n")

print("任务完成...")
f.close()